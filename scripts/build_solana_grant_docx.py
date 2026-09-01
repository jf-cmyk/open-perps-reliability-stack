from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "Open Perps Reliability Stack - Solana Foundation Proposal.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "EAF3F8"
LIGHT_GRAY = "F4F6F9"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for row in rows:
        table_row = table.add_row()
        set_row_cant_split(table_row)
        cells = table_row.cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    document.add_paragraph()
    return table


def add_callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(10)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(10)
    document.add_paragraph()


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_bullet(document, text):
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f" {text}")


def add_number(document, text):
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.194)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Open Perps Reliability Stack For Solana")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle_run = subtitle.add_run("Solana Foundation Developer Tooling Grant Proposal")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = DARK_BLUE

    add_table(
        document,
        ["Field", "Response"],
        [
            ["Applicant", "Blocksize"],
            ["Project", "Open Perps Reliability Stack for Solana"],
            ["Funding category", "Developer Tooling"],
            ["Funding request", "$125,000"],
            ["Public repo", "https://github.com/jf-cmyk/open-perps-reliability-stack"],
            ["Scope", "Open-source, read-only, dry-run only"],
        ],
        [2200, 7160],
    )

    add_callout(
        document,
        "Grant scope boundary",
        "The grant-funded scope is intentionally read-only and dry-run only. It does not include signing, custody, capital deployment, live liquidation execution, live transaction submission, private-key handling, or trading-profit promises.",
    )

    document.add_heading("Summary", level=1)
    document.add_paragraph(
        "Open Perps Reliability Stack is open-source developer tooling for Solana onchain perps reliability. It provides read-only protocol adapters, source-governed decode records, Pyth-aware risk primitives, normalized market-quality data, fixture-backed deterministic dry-run reasoning, and public dashboard/API schemas."
    )
    document.add_paragraph(
        "The current MVP is intentionally complementary infrastructure, not a venue. It gives builders and reviewers a reproducible way to inspect reliability evidence across Drift, Jupiter Perps, Phoenix/Rise, oracle inputs, and public proof packages while preserving a strict no-signing and no-production-execution boundary."
    )

    document.add_heading("Problem", level=1)
    document.add_paragraph(
        "Solana perps venues are becoming more important, but reliability tooling is fragmented. Builders and risk teams need to understand market quality, oracle health, liquidation conditions, adapter/decode quality, and dry-run execution-failure methodology under stress. Today these views are often venue-specific, private, or embedded inside searcher infrastructure."
    )
    for item in [
        "Which oracle feeds are stale, wide, missing, or diverging from venue marks?",
        "Which markets show weak depth, high funding stress, or poor liquidation health?",
        "Which account/event decoders are failing after program or IDL changes?",
        "Which liquidation windows are detectable from public state, and why would a dry-run accept or reject them?",
        "Which data can be published safely without leaking private strategy, RPC secrets, or signer/custody details?",
    ]:
        add_bullet(document, item)

    document.add_heading("Proposed Solution", level=1)
    for item in [
        "Adapter standard for read-only Solana perps venue integrations.",
        "First fixture-backed Drift read-only adapter.",
        "Source-review records that gate binary decode and replay claims on canonical source authority.",
        "Pyth-aware risk SDK for staleness, confidence, divergence, and liquidation-risk inputs.",
        "Canonical event envelope, lineage model, dataset manifests, and data quality publish gates.",
        "Deterministic replay fixtures and dry-run output bundles with explicit reason codes.",
        "Public dashboard/API schemas for market quality, oracle risk, liquidation health, adapter health, and methodology.",
    ]:
        add_bullet(document, item)
    document.add_paragraph(
        "This is not a new perp venue. It is neutral tooling that helps existing and future Solana perps systems become more inspectable, testable, and reliable."
    )

    document.add_heading("Why Solana", level=1)
    for item in [
        "Pyth confidence, exponent, publish time, and freshness.",
        "Slot timing, blockhash expiry, priority fees, compute budget, and account locks.",
        "Venue-specific program accounts and event layouts.",
        "High-throughput historical state, replay, and public dataset needs.",
        "Onchain liquidation and margin semantics that differ by Solana venue.",
        "Fast-moving network regimes such as VAT, Alpenglow, Agave, Firedancer, and DoubleZero transport changes.",
    ]:
        add_bullet(document, item)

    document.add_heading("2026 Ecosystem Refresh And Grant Positioning", level=1)
    document.add_paragraph(
        "The latest source-backed research strengthens the grant case while narrowing what should be claimed. Solana Foundation's 2026 Open Perps call explicitly welcomes complementary infrastructure around fully onchain perps; OPRS should therefore position the public-good core as reliability infrastructure rather than as execution, trading, or liquidation operations."
    )
    for item in [
        "Drift remains the first fixture-backed adapter path. Historical liquidation reconstruction should advance only from public finalized transaction evidence and pinned legacy source, with migrated Velocity-hosted records used as discovery or corroboration rather than sole authority.",
        "Jupiter Perps is relevant and should stay in scope, but its canonical position is source-authority blocked: first-party docs support lifecycle semantics, while binary account decoding and deterministic request-to-position pairing require a current Jupiter-confirmed IDL/source or hashable artifact.",
        "Phoenix/Rise is now a source-pinned follow-on venue lane because public gold and crude-oil perps surfaces expose mark, index, volume, open interest, funding, and market state. Production program and Hawkeye view constants are source-pinned from Ellipsis Labs Rise, and a Phoenix/Hawkeye validator-plan contract records scrubbed fixture gates, while exact oracle/input identities, account-level decode, trader monitoring, and replay remain blocked.",
        "Frontier Traders creates a review-only design-partner channel for professional trader reliability feedback, but it does not establish Blocksize access, demand, or endorsement.",
        "Pay.sh, Solana Subscriptions, and Commerce Kit are useful for future commercial packaging and revenue-routing clarity. They should stay outside the grant-funded public-good core unless explicitly framed as optional or convertible-grant scope.",
        "The 100M-CU block-capacity activation and the source-governed 400ms-to-350ms slot-time activation at slot 440208000 add new benchmark regimes for aggregate block headroom, slot-normalized landing, unchanged hot-account contention, replay, catchup, and downstream infrastructure behavior. They are context for reliability methodology, not evidence of measured performance improvement.",
        "Network changes around VAT, Alpenglow, Agave, Firedancer, and DoubleZero make validator/client/transport context relevant to perps reliability, but releases, schedules, and gossip records do not prove Blocksize adoption or measured performance gains.",
    ]:
        add_bullet(document, item)

    document.add_page_break()
    document.add_heading("Milestones And Budget", level=1)
    add_table(
        document,
        ["Milestone", "Funding", "Core Deliverables", "Evidence"],
        [
            [
                "1. Adapter and architecture foundation",
                "$25,000",
                "Architecture v0, ADRs, adapter standard, capability matrix, Rust workspace, fixture-backed Drift read-only adapter, source-authority records.",
                "Public repo, compile-passing workspace, adapter metadata, execute-disabled safety tests, source-review validators.",
            ],
            [
                "2. Canonical data and sample datasets",
                "$35,000",
                "Canonical event envelope, dataset manifest, DQ gates, scrub policy, golden fixtures, Drift liquidation-history probe contract.",
                "Rust schema types, JSON examples, checksum-bearing fixture manifest, scrubbed discovery-output schema.",
            ],
            [
                "3. Pyth-aware risk SDK and dry-run replay",
                "$40,000",
                "Oracle staleness/confidence/divergence primitives, liquidation state, dry-run outputs, reason taxonomy.",
                "Unit tests, dry-run samples, no-signing transaction-plan guardrails.",
            ],
            [
                "4. Public dashboard/API contract and final report",
                "$25,000",
                "Public API schema, dashboard view contract, demo narrative, source-governed methodology, limitations, and venue-readiness matrix.",
                "Schema files, sample responses, final report, reproducible demo path, Railway and GitHub Pages proof links.",
            ],
        ],
        [1900, 1100, 3350, 3010],
    )

    document.add_heading("Product And Adoption Metrics", level=1)
    document.add_heading("Developer-tooling metrics", level=2)
    for item in [
        "GitHub stars, forks, issues, and adapter requests from Solana builders and researchers.",
        "Adapter standard reuse or comments by perps and infrastructure teams.",
        "Number of public fixture runs and replay examples.",
        "Number of supported markets and venues in read-only mode.",
        "Public dataset downloads or downstream notebooks.",
        "Dashboard/API schema consumers.",
        "Source-authority confirmations, issue comments, or protocol maintainer reviews that unblock adapter promotion.",
    ]:
        add_bullet(document, item)
    document.add_heading("Technical quality metrics", level=2)
    for item in [
        "Adapter decode success rate and schema mismatch rate.",
        "Oracle stale, wide, missing, or divergent feed counts.",
        "Dry-run reason-code distribution.",
        "Fixture replay pass/fail status.",
        "Dataset publish gate pass/warn/block status.",
        "Source-review approval status by account type, program ID, commit/hash, and public regression fixture availability.",
    ]:
        add_bullet(document, item)

    document.add_heading("Public-Good Value", level=1)
    document.add_paragraph(
        "The project turns private reliability knowledge into reusable Solana developer tooling: open adapters, open risk primitives, open data schemas, open replay and dry-run contracts, public sample datasets, public dashboard/API methodology, and clear limitations."
    )

    document.add_heading("Commercial Boundary", level=1)
    document.add_paragraph(
        "Blocksize may later build commercial services around managed integrations, premium APIs, private analytics, and controlled execution tooling. Those are outside this grant. Grant-funded work remains public, reproducible, read-only, and no-signing. Commercial work cannot privatize the OSS artifacts funded by the grant."
    )

    document.add_heading("Current Proof Of Work", level=1)
    for item in [
        "OSS hygiene files and GitHub workflow.",
        "Architecture docs, roadmap, work packages, and ADRs.",
        "Rust workspace with core crates.",
        "Fixture-backed Drift read-only adapter.",
        "Adapter metadata and capability model.",
        "Canonical event and dataset manifest types.",
        "Data quality publish gate and scrub policy types.",
        "Pyth-aware risk primitives with unit tests.",
        "Replay fixture and dry-run output contracts.",
        "Data reconstruction envelope schema with provider, commitment, slot range, evidence references, known gaps, and scrub-policy validation.",
        "Expanded Solana runtime failure reason codes for deterministic dry-run explanation.",
        "Read-only target discovery command for Helius-backed Drift/Jupiter proof setup; first local run now succeeds and writes scrubbed output under target/.",
        "Jupiter source-authority confirmation packet and send-ready outbound note.",
        "Source-review record schema and examples for Jupiter authority confirmation and Drift public-field promotion.",
        "Scrubbed Drift liquidation-history probe schema, validator, and bounded discovery output contract.",
        "Drift legacy liquidation-history diligence has scanned 288,000 finalized program transactions from July 22 back through slot 415592674 at 2026-04-25T15:31:16Z without a matching Liquidate* log; this is queue progress only, not evidence that liquidations were absent.",
        "Slot-regime benchmark package for the source-governed 400ms-to-350ms activation boundary at slot 440208000, with pre/post windows for future read-only normalization and explicit no-performance-claim gates.",
        "Continuous Solana ecosystem research loop with hot state, evidence ledger, opportunity pipeline, and checkpoint archive.",
        "Hosted Railway proof-pack MVP and dashboard with filtered GitHub Pages fallback.",
        "Hourly hosted smoke monitoring and public artifact boundary checks.",
        "Grant package and application draft.",
    ]:
        add_bullet(document, item)

    add_callout(
        document,
        "Reviewer proof pack",
        "Canonical proof pack: https://refreshing-art-production-86de.up.railway.app/. Dashboard: https://refreshing-art-production-86de.up.railway.app/apps/dashboard/. GitHub Pages remains an equivalent filtered fallback. The MVP proof checklist maps claims to URLs, schemas, fixtures, and validation commands in docs/mvp-proof-checklist.md.",
    )

    document.add_heading("Why Blocksize", level=1)
    document.add_paragraph(
        "Blocksize brings Solana infrastructure experience, validator/reliability context, Pyth data-provider experience, and prior liquidation/reliability research. The team understands the difference between public-good methodology and private execution alpha, which is important for keeping this grant clean."
    )
    document.add_paragraph(
        "The project benefits from a disciplined scope boundary: no signing, no custody, no production execution, and no capital deployment in the grant phase. That makes the first phase useful to the ecosystem while avoiding the security, market-conduct, and compliance risks of live liquidation operations."
    )

    document.add_heading("Risks And Mitigations", level=1)
    risks = [
        ("Perps-specific historical liquidation data is thin.", "Start with fixture-backed Drift adapter shape tests, synthetic golden fixtures, and explicit data-quality caveats before claiming replay coverage."),
        ("Venue schemas drift.", "Adapter metadata includes program IDs, schema versions, supported account schema versions, IDL hash, source update timestamps, and caveats."),
        ("Jupiter source authority is not yet canonical.", "Use first-party Jupiter docs only for semantic labels. Keep binary decode, deterministic request/fulfillment pairing, and historical replay blocked until Jupiter provides or confirms a current hashable IDL/source and fixtures."),
        ("Phoenix public telemetry is not yet account-level replay proof.", "Use public GOLD and external-asset surfaces plus source-pinned program/Hawkeye constants for adapter-contract design while keeping exact oracle inputs, account layouts, and transaction fixtures as explicit gates."),
        ("Research findings could drift into overclaiming.", "Maintain the Solana ecosystem loop as source-backed project memory, separate fact from inference, and require proof artifacts before promoting any partner, demand, revenue, or protocol-safety claim."),
        ("Public datasets leak private information.", "Publish gates and scrub policy remove RPC URLs, API keys, internal paths, route labels, private strategy thresholds, capital controls, and signer/custody metadata."),
        ("Historical Helius decode proof is not complete yet.", "The local read-only target discovery command succeeds, but deeper Drift market/oracle and Jupiter pool/custody decode coverage still needs public target resolution and scrubbed proof output before it can be claimed."),
        ("Project scope creeps into execution.", "Dry-run transaction plans require requires_signer=false and submission_disabled=true; production execution remains out of scope."),
        ("Commercial track could confuse grant reviewers.", "Grant-funded outputs are explicitly public-good OSS. Commercial services are disclosed as future/out-of-scope and cannot privatize grant-funded modules."),
    ]
    for idx, (risk, mitigation) in enumerate(risks, start=1):
        add_number(document, f"{risk} Mitigation: {mitigation}")

    document.add_heading("Submission Notes", level=1)
    for item in [
        "Recommended funding category: Developer Tooling.",
        "Recommended form amount: $125,000.",
        "Recommended on-chain accounts field: N/A - read-only and dry-run only.",
        "Recommended project/idea field: short summary plus this proposal document and public GitHub repo.",
    ]:
        add_bullet(document, item)

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
